"""Rellena las plantillas Word fijas (GTH) con datos del trabajador y del plan."""

from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass
from datetime import date
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table

from .calendar import DERECHO_ANUAL, parse_iso_date, vacation_record_for
from ..textnorm import strip_marks

TEMPLATES_DIR = Path(__file__).resolve().parents[1] / "data" / "templates"

TEMPLATES = {
    1: ("escenario_1_memorando.docx", "Memorando_de_vacaciones"),
    2: ("escenario_2_fraccionamiento.docx", "Fraccionamiento_vacacional"),
    3: ("escenario_3_modificacion.docx", "Modificacion_fraccionamiento"),
    4: ("escenario_4_adelanto.docx", "Adelanto_de_vacaciones"),
}

TITULOS = {
    1: "Memorando de vacaciones",
    2: "Fraccionamiento de descanso vacacional",
    3: "Modificación de fraccionamiento",
    4: "Adelanto de goce de vacaciones",
}

# Compacto (sin espacios) después de _norm_empresa. II va primero.
EMPRESAS: tuple[tuple[str, str, frozenset[str]], ...] = (
    ("AQU ANQA II S.A.C.", "20610068767", frozenset({"AQUII", "AQUANQAII"})),
    ("AQU ANQA S.A.C.", "20608345770", frozenset({"AQU", "AQUANQA"})),
)

MESES = (
    "",
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
)

_NS_XML = "{http://www.w3.org/XML/1998/namespace}space"
_PLACEHOLDER = r"(?:x+|X+|…+)"
_FLAGS = re.IGNORECASE


@dataclass(frozen=True)
class DocContext:
    fecha: date
    nombre: str
    dni: str
    empresa: str
    ruc: str
    jefe: str
    cargo_jefe: str
    record: str
    inicio: date
    fin: date
    dias: int
    derecho_desde: date | None
    periodos: tuple[dict, ...]
    periodos_anteriores: tuple[dict, ...]


def _norm_empresa(empresa: str) -> str:
    raw = strip_marks(empresa).upper().replace(".", " ")
    raw = re.sub(r"[^A-Z0-9]+", " ", raw)
    raw = re.sub(r"\bS\s*A\s*C\b", "", raw)
    return re.sub(r"\s+", " ", raw).strip()


def empresa_legal(empresa: str) -> tuple[str, str]:
    compact = _norm_empresa(empresa).replace(" ", "")
    for razon, ruc, aliases in EMPRESAS:
        if compact in aliases:
            return razon, ruc
    return (empresa or "").strip(), ""


def infer_escenario(*, es_adelanto: bool, moved: bool, period_sizes: list[int]) -> int:
    if moved:
        return 3
    if es_adelanto:
        return 4
    if len(period_sizes) == 1 and period_sizes[0] == DERECHO_ANUAL:
        return 1
    return 2


def documento_meta(*, es_adelanto: bool, moved: bool, period_sizes: list[int]) -> dict:
    escenario = infer_escenario(es_adelanto=es_adelanto, moved=moved, period_sizes=period_sizes)
    return {"escenario": escenario, "titulo": TITULOS[escenario]}


def fecha_larga(d: date) -> str:
    return f"{d.day} de {MESES[d.month]} de {d.year}"


def fecha_slash(d: date) -> str:
    return d.strftime("%d/%m/%Y")


def rango_narrativo(inicio: date, fin: date) -> str:
    if inicio.year == fin.year and inicio.month == fin.month:
        return f"del {inicio.day} al {fin.day} de {MESES[inicio.month]} del año {inicio.year}"
    if inicio.year == fin.year:
        return (
            f"del {inicio.day} de {MESES[inicio.month]} al {fin.day} de {MESES[fin.month]} "
            f"del año {inicio.year}"
        )
    return (
        f"del {inicio.day} de {MESES[inicio.month]} de {inicio.year} "
        f"al {fin.day} de {MESES[fin.month]} de {fin.year}"
    )


def _direct_xml_text(p_el) -> str:
    parts: list[str] = []
    for child in p_el:
        if child.tag == qn("w:r"):
            if child.find(qn("w:drawing")) is not None or child.find(qn("w:pict")) is not None:
                continue
            for t in child.findall(qn("w:t")):
                parts.append(t.text or "")
        elif child.tag in (qn("w:hyperlink"), qn("w:ins"), qn("w:smartTag")):
            for t in child.findall(f".//{qn('w:t')}"):
                parts.append(t.text or "")
    return "".join(parts)


def _set_xml_paragraph_text(p_el, text: str) -> None:
    text_nodes: list = []
    for child in p_el:
        if child.tag != qn("w:r"):
            continue
        if child.find(qn("w:drawing")) is not None or child.find(qn("w:pict")) is not None:
            continue
        text_nodes.extend(child.findall(qn("w:t")))
    if text_nodes:
        text_nodes[0].text = text
        text_nodes[0].set(_NS_XML, "preserve")
        for node in text_nodes[1:]:
            node.text = ""
        return
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.set(_NS_XML, "preserve")
    node.text = text
    run.append(node)
    p_el.append(run)


def _is_blank_body_para(p_el) -> bool:
    if _direct_xml_text(p_el).strip():
        return False
    if p_el.find(f".//{qn('w:drawing')}") is not None:
        return False
    if p_el.find(f".//{qn('w:pict')}") is not None:
        return False
    return True


def _tighten_blank_paragraphs(doc: Document, keep: int = 1) -> None:
    body = doc.element.body
    streak: list = []
    to_remove: list = []

    def flush() -> None:
        if len(streak) > keep:
            to_remove.extend(streak[keep:])
        streak.clear()

    for child in list(body):
        if child.tag == qn("w:p") and _is_blank_body_para(child):
            streak.append(child)
        else:
            flush()
    flush()
    for el in to_remove:
        body.remove(el)


def _fill_whole_line(stripped: str, ctx: DocContext) -> str | None:
    """Párrafo que es solo una etiqueta; se reemplaza entero."""
    if re.fullmatch(r"Fecha[:;]\s*(x+)?", stripped, _FLAGS):
        return f"Fecha: {fecha_larga(ctx.fecha)}"
    if re.fullmatch(r"Se[ñn]or\s*\(Sra\.?\)\s*:?", stripped, _FLAGS):
        return f"Señor (Sra.): {ctx.nombre}"
    if (
        re.search(r"nombres?\s+y\s+apellidos", stripped, _FLAGS)
        and re.search(r"\bdni\b", stripped, _FLAGS)
        and len(stripped) < 90
    ):
        return f"{ctx.nombre}\nDNI N°: {ctx.dni}"
    if re.fullmatch(r"nombres y apellidos( del trabajador)?", stripped, _FLAGS):
        return ctx.nombre
    if re.fullmatch(r"nombre y apellidos", stripped, _FLAGS):
        return ctx.nombre
    if re.fullmatch(r"DNI(\s*N[°º])?:?", stripped, _FLAGS):
        return f"DNI N°: {ctx.dni}"
    if stripped.casefold() == "nombre del jefe de sub área o jefe inmediato" and ctx.jefe:
        return ctx.jefe
    if re.fullmatch(r"\(Cargo\)(\s*x+)?", stripped, _FLAGS):
        return ctx.cargo_jefe or "(Cargo)"
    return None


def _inline_subs(ctx: DocContext) -> list[tuple[str, str]]:
    """Orden fijo: días antes que récord, para no dejar 'POR xxxx DÍAS'."""
    ruc = ctx.ruc or "________________"
    empresa = ctx.empresa or "________________"
    firma_fecha = (
        f"a los {ctx.fecha.day} días del mes de {MESES[ctx.fecha.month]} del año {ctx.fecha.year}"
    )
    goce = (
        f"programándose del {fecha_slash(ctx.inicio)} al {fecha_slash(ctx.fin)} "
        f"del año {ctx.fin.year}"
    )
    rows: list[tuple[str, str]] = [
        (r"\bPOR\s+x+\s+D[ÍI]AS\b", f"POR {ctx.dias} DÍAS"),
        (r"\(COLOCAR LA RAZÓN SOCIAL EN LA QUE ESTÉ CONTRATADO\)", empresa),
        (r"\(COLOCAR RUC\)", ruc),
        (r"\(RAZ[OÓ]N SOCIAL\)", empresa),
        (r"con RUC N° \(COLOCAR RUC\)", f"con RUC N° {ruc}"),
        (r"con RUC\s+" + _PLACEHOLDER, f"con RUC {ruc}"),
        (r"El \(A\) Señor \(Sra\)\s+" + _PLACEHOLDER, f"El (A) Señor (Sra) {ctx.nombre}"),
        (r"el Sr \(a\)\.\s*" + _PLACEHOLDER, f"el Sr (a). {ctx.nombre}"),
        (r"identificado con DNI N[°º]\s*" + _PLACEHOLDER, f"identificado con DNI N° {ctx.dni}"),
        (r"con DNI N[°º]\s*" + _PLACEHOLDER, f"con DNI N° {ctx.dni}"),
        (rf"(RECORD VACACIONAL)\s+{_PLACEHOLDER}", rf"\1 {ctx.record}"),
        (rf"(r[ée]cord vacacional)\s+{_PLACEHOLDER}", rf"\1 {ctx.record}"),
        (rf"(del r[ée]cord)\s+{_PLACEHOLDER}", rf"\1 {ctx.record}"),
        (r"r[ée]cord 2025-2026", f"récord {ctx.record}"),
        (r"programándose del x+ al x+ del año x+", goce),
        (r"del xxx al x+ de junio del año x+", rango_narrativo(ctx.inicio, ctx.fin)),
        (
            r"del\s+x+\s+al\s+x+\s+de\s+x+\s+del\s+(?:a[ñn]o\s+)?" + _PLACEHOLDER,
            rango_narrativo(ctx.inicio, ctx.fin),
        ),
        (r"por xx días, del xx al xx de xxx del x+", f"por {ctx.dias} días, {rango_narrativo(ctx.inicio, ctx.fin)}"),
        (r"a los x+ días del mes de x+ del año x+", firma_fecha),
        (r"por XXX días", f"por {ctx.dias} días"),
        (r"otorga X+\s+días", f"otorga {ctx.dias} días"),
        (r"Con fecha x+", f"Con fecha {fecha_larga(ctx.fecha)}"),
    ]
    if ctx.derecho_desde:
        rows.append(
            (r"a partir del día " + _PLACEHOLDER, f"a partir del día {fecha_slash(ctx.derecho_desde)}")
        )
    return rows


def fill_text(text: str, ctx: DocContext) -> str:
    whole = _fill_whole_line(text.strip(), ctx)
    if whole is not None:
        return whole
    filled = text
    for pattern, replacement in _inline_subs(ctx):
        filled = re.sub(pattern, replacement, filled, flags=_FLAGS)
    return filled


def _fill_periods_table(table: Table, periods: list[dict] | tuple[dict, ...]) -> None:
    labels = ("Primer periodo", "Segundo periodo", "Tercer periodo", "Cuarto periodo", "Quinto periodo")
    needed = max(len(periods), 3)
    while len(table.rows) < needed + 1:
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    for i in range(needed):
        row = table.rows[i + 1]
        label = labels[i] if i < len(labels) else f"Periodo {i + 1}"
        if i < len(periods):
            p = periods[i]
            vals = (label, str(p["dias"]), fecha_slash(p["inicio"]), fecha_slash(p["fin"]))
        else:
            vals = (label, "", "", "")
        for cell, val in zip(row.cells, vals):
            if cell.paragraphs:
                _set_xml_paragraph_text(cell.paragraphs[0]._p, val)
            else:
                cell.text = val


def _period_sets_for_tables(escenario: int, ctx: DocContext, n_tables: int) -> list[tuple[dict, ...]]:
    if escenario == 2:
        return [ctx.periodos] * n_tables
    if escenario == 3:
        anteriores = ctx.periodos_anteriores or ctx.periodos
        layout = [ctx.periodos, anteriores, ctx.periodos]
        return layout[:n_tables]
    return []


def build_context(
    emp: dict,
    *,
    today: date,
    year: int,
    inicio: date,
    fin: date,
    dias: int,
    periodos: list[dict],
    periodos_anteriores: list[dict] | None = None,
    programmed: list[date] | None = None,
) -> DocContext:
    ingreso = parse_iso_date(emp.get("fecha_ingreso"))
    rec = vacation_record_for(ingreso, programmed or [], year, today)
    jefatura = (emp.get("jefatura") or "").strip()
    gerencia = (emp.get("gerencia") or "").strip()
    razon, ruc = empresa_legal((emp.get("empresa") or "").strip())
    cumple = rec.get("cumple_record")
    return DocContext(
        fecha=today,
        nombre=(emp.get("nombre") or "").strip(),
        dni=str(emp.get("dni") or "").strip(),
        empresa=razon,
        ruc=ruc,
        jefe=jefatura,
        cargo_jefe=gerencia or (f"Jefe de {jefatura}" if jefatura else ""),
        record=rec.get("record_vacacional") or f"{year - 1}-{year}",
        inicio=inicio,
        fin=fin,
        dias=dias,
        derecho_desde=cumple if isinstance(cumple, date) else None,
        periodos=tuple(periodos),
        periodos_anteriores=tuple(periodos_anteriores or []),
    )


def fill_template(escenario: int, ctx: DocContext) -> bytes:
    if escenario not in TEMPLATES:
        raise ValueError("No hay plantilla para ese escenario.")
    name, _slug = TEMPLATES[escenario]
    path = TEMPLATES_DIR / name
    if not path.is_file():
        raise FileNotFoundError(f"Falta la plantilla {name} en {TEMPLATES_DIR}.")

    doc = Document(str(path))
    for p_el in doc.element.body.iter(qn("w:p")):
        original = _direct_xml_text(p_el)
        if not original.strip():
            continue
        filled = fill_text(original, ctx)
        if filled != original:
            _set_xml_paragraph_text(p_el, filled)

    tables = list(doc.tables)
    for table, periods in zip(tables, _period_sets_for_tables(escenario, ctx, len(tables))):
        _fill_periods_table(table, periods)

    _tighten_blank_paragraphs(doc, keep=1)
    _strip_comment_markup(doc.element)
    buf = BytesIO()
    doc.save(buf)
    return _strip_comment_parts(buf.getvalue())


_COMMENT_PARTS = (
    "word/comments.xml",
    "word/commentsExtended.xml",
    "word/commentsIds.xml",
    "word/commentsExtensible.xml",
    "word/people.xml",
)
_COMMENT_NS_TAGS = (
    qn("w:commentRangeStart"),
    qn("w:commentRangeEnd"),
    qn("w:commentReference"),
)


def _strip_comment_markup(root) -> None:
    """Quita marcas de comentario de Word (globos / panel de revisión)."""
    for tag in _COMMENT_NS_TAGS:
        for node in root.findall(f".//{tag}"):
            parent = node.getparent()
            if parent is None:
                continue
            parent.remove(node)
            if parent.tag == qn("w:r") and not any(c.tag != qn("w:rPr") for c in parent):
                grand = parent.getparent()
                if grand is not None:
                    grand.remove(parent)


def _drop_xml_nodes(xml: bytes, *, attr: str, needles: tuple[str, ...]) -> bytes:
    text = xml.decode("utf-8")
    pattern = (
        r"<[^>]*"
        + re.escape(attr)
        + r'="[^"]*(?:'
        + "|".join(re.escape(n) for n in needles)
        + r')[^"]*"[^>]*/>'
    )
    return re.sub(pattern, "", text).encode("utf-8")


def _strip_comment_parts(data: bytes) -> bytes:
    """Elimina partes de comentarios del paquete .docx."""
    src = ZipFile(BytesIO(data))
    out = BytesIO()
    drop = set(_COMMENT_PARTS)
    with ZipFile(out, "w", ZIP_DEFLATED) as dest:
        for info in src.infolist():
            if info.filename in drop:
                continue
            payload = src.read(info.filename)
            if info.filename == "[Content_Types].xml":
                payload = _drop_xml_nodes(
                    payload, attr="PartName", needles=("/word/comments", "/word/people.xml")
                )
            elif info.filename.endswith(".rels"):
                payload = _drop_xml_nodes(
                    payload, attr="Type", needles=("relationships/comments", "relationships/people")
                )
            dest.writestr(info, payload)
    src.close()
    return out.getvalue()


def filename_for(escenario: int, ctx: DocContext) -> str:
    _name, slug = TEMPLATES[escenario]
    safe_dni = re.sub(r"[^0-9A-Za-z]", "", ctx.dni) or "trabajador"
    return f"{slug}_{safe_dni}_{ctx.inicio.isoformat()}.docx"


def reconstruct_old_periods(
    new_periods: list[dict], old_start: date, new_start: date
) -> list[dict]:
    out: list[dict] = []
    for p in new_periods:
        if p["inicio"] == new_start:
            delta = p["fin"] - p["inicio"]
            out.append({"inicio": old_start, "fin": old_start + delta, "dias": p["dias"]})
        else:
            out.append(dict(p))
    out.sort(key=lambda x: x["inicio"])
    return out
